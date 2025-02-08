from docx import Document
from docx.shared import Inches
import os

def pDocx(ClientName, ClientID, current_date, response, serviceProvided, serviceProvidedBy, startTime, endTime, totaltime, timeQuarter):
    document = Document()
    p = document.add_paragraph('Service Log' + '\n' + current_date.strftime('%m-%d-%Y') + "\n")
    p.add_run(ClientName + "\n")
    p.add_run('Medicaid Number: ' + ClientID + "\n")
    p.add_run(serviceProvided + "\n")
    p.add_run(serviceProvidedBy + "\n")
    p.add_run('Peterson Family Care LLC Provider Number: 009279700' + "\n")
    p.add_run('Start Time: ' + str(startTime) + ' End Time: ' + str(endTime) + "\n")
    p.add_run('Total Hours: ' + str(totaltime) + "\n")
    p.add_run('Total Quarter Hours: ' + str(timeQuarter) + "\n")
    
    document.add_heading('Daily Report', level=1)
    document.add_paragraph(response.message.content + "\n")

    if not os.path.exists(ClientName):
        os.mkdir(ClientName)
    
    document.save(str(ClientName) + '/' + str(ClientName).replace(" ", "-") + ' ' + current_date.strftime('%m-%d-%Y') + ".docx")